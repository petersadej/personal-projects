import pandas as pd
import sys
import re
import statistics
from pathlib import Path
from docx import Document
from docx.shared import Pt


def read_input_file(input_file: str) -> pd.DataFrame:
    """
    Read supported tabular input files (Excel or CSV) into a DataFrame.

    Args:
        input_file: Path to the input file

    Returns:
        pandas.DataFrame
    """
    file_extension = Path(input_file).suffix.lower()

    if file_extension == '.csv':
        return pd.read_csv(input_file)

    if file_extension in ['.xlsx', '.xls', '.xlsm']:
        return pd.read_excel(input_file)

    raise ValueError(
        f"Unsupported file type '{file_extension}'. Supported types: .csv, .xlsx, .xls, .xlsm"
    )


def excel_to_word(excel_file: str, output_file: str = None) -> str:
    """
    Convert an Excel file to a formatted Word document.
    Each row is written with column headers in bold preceding the data.

    Args:
        excel_file: Path to the Excel file
        output_file: Path to the output Word file (optional)
    """
    try:
        # Read the input file
        df = read_input_file(excel_file)

        # If no output file specified, create one based on input filename
        if output_file is None:
            output_file = Path(excel_file).stem + "_output.docx"

        # Create a new Word document
        doc = Document()

        # Get column headers
        headers = df.columns.tolist()

        # Process each row
        for index, row in df.iterrows():
            # Add row header
            row_heading = doc.add_paragraph()
            run = row_heading.add_run(f"Row {index + 1}:")
            run.bold = True
            run.font.size = Pt(12)

            # Add separator
            doc.add_paragraph("_" * 50)

            # Write each column header with its corresponding value
            for header in headers:
                value = row[header]
                p = doc.add_paragraph()
                # Add header in bold
                header_run = p.add_run(f"{header}: ")
                header_run.bold = True
                # Add value in normal text
                p.add_run(str(value))

            # Add blank paragraph between rows
            doc.add_paragraph()

        # Save the document
        doc.save(output_file)

        print(f"Successfully converted {excel_file} to {output_file}")
        return output_file

    except FileNotFoundError:
        print(f"Error: File '{excel_file}' not found.")
        sys.exit(1)
    except Exception as e:
        print(f"Error: {e}")
        sys.exit(1)


def specific_excel_to_word(excel_file: str, output_file: str = None) -> str:
    """
    Convert a specific Excel layout to a formatted Word document.

    Special parsing rules:


    Args:
        excel_file: Path to the Excel file
        output_file: Path to the output Word file (optional)
    """
    try:
        def parse_likert_answer(value):
            if pd.isna(value):
                return None

            value_text = str(value).strip()
            normalized_value_text = re.sub(r'\s+', ' ', value_text.lower()).strip()

            word_scale_map = {
                'not at all confident': 1,
                'slightly confident': 2,
                'moderately confident': 3,
                'very confident': 4,
                'extremely confident': 5,
                'extrmely confident': 5,
            }

            if normalized_value_text in word_scale_map:
                return word_scale_map[normalized_value_text]

            # Accept direct numeric answers 1-5
            if re.fullmatch(r'[1-5]', value_text):
                return int(value_text)

            # Accept labels like "1 = No more needed"
            match = re.match(r'^([1-5])\s*=', value_text)
            if match:
                return int(match.group(1))

            return None

        def is_likert_column(header, dataframe):
            cleaned_header = re.sub(r'^(\s*\d+\.)+\s*', '', str(header)).strip().lower()
            if re.search(r'\b1\s*=.*\b5\s*=', cleaned_header):
                return True

            column_values = dataframe[header].dropna()
            non_empty_values = [str(value).strip() for value in column_values if str(value).strip()]
            if not non_empty_values:
                return False

            parsed_count = sum(1 for value in non_empty_values if parse_likert_answer(value) is not None)
            # Treat as Likert when most non-empty values map to the 1-5 scale
            return parsed_count >= max(1, int(len(non_empty_values) * 0.7))

        def preprocess_likert_columns(dataframe, likert_column_headers):
            for header in likert_column_headers:
                dataframe[header] = dataframe[header].apply(parse_likert_answer)

            return dataframe

        # Read the input file
        df = read_input_file(excel_file)

        if output_file is None:
            output_file = Path(excel_file).stem + "_special_output.docx"

        # Create a new Word document
        doc = Document()

        # Get column headers
        headers = df.columns.tolist()

        # Analyze special-purpose headers in a single pass
        jurisdiction_column = None
        pre_symposium_headers = []
        other_specify_headers = []
        sww_impacts_project_headers = []
        sww_impacts_project_labels = {}
        success_story_headers = {}
        likert_headers = []
        participant_headers = {}
        ambiguous_participant_index = 1

        for header in headers:
            cleaned_header = re.sub(r'^(\s*\d+\.)+\s*', '', str(header)).strip()
            cleaned_header_lower = cleaned_header.lower()

            if jurisdiction_column is None and 'jurisdiction' in cleaned_header_lower:
                jurisdiction_column = header

            if 'pre-symposium virtual learning collaborative sessions will occur on' in cleaned_header_lower:
                pre_symposium_headers.append(header)

            if "you selected 'other'. please specify" in cleaned_header_lower:
                other_specify_headers.append(header)

            if 'which sww-impacts project will you focus on?' in cleaned_header_lower:
                sww_impacts_project_headers.append(header)
                choice_match = re.search(r'choice=<b>(.*?)</b>', str(header), flags=re.IGNORECASE)
                if choice_match:
                    sww_impacts_project_labels[header] = choice_match.group(1).strip()
                else:
                    sww_impacts_project_labels[header] = cleaned_header

            if cleaned_header_lower.startswith(
                'interventions with children or adolescents to decrease cancer risk'
            ):
                success_story_headers[header] = 'Addressing Risk Factors'

            if cleaned_header_lower.startswith(
                'health systems changes to address cancer risk factors'
            ):
                success_story_headers[header] = 'Health systems'

            if is_likert_column(header, df):
                likert_headers.append(header)

            participant_match = re.search(r'participant\s+(\d+)', cleaned_header, flags=re.IGNORECASE)

            if participant_match:
                participant_number = int(participant_match.group(1))
                participant_headers.setdefault(participant_number, []).append(header)
                continue

            # Map repeated "Did this person..." columns to participants by position
            if 'did this person participate in the sww 2024 pilot program?' in cleaned_header.lower():
                participant_headers.setdefault(ambiguous_participant_index, []).append(header)
                ambiguous_participant_index += 1

        participant_header_set = {
            header
            for grouped_headers in participant_headers.values()
            for header in grouped_headers
        }

        df = preprocess_likert_columns(df, likert_headers)

        likert_values_by_header = {header: [] for header in likert_headers}

        # Process each row
        for index, row in df.iterrows():
            # Add row header
            row_heading = doc.add_paragraph()

            program_value = ""
            if jurisdiction_column is not None and pd.notna(row[jurisdiction_column]):
                program_value = str(row[jurisdiction_column]).strip()

            run = row_heading.add_run(f"Program: {program_value}")
            run.bold = True
            run.font.size = Pt(16)

            # Add separator
            doc.add_paragraph("_" * 50)

            # Write participant data grouped into participant sections first
            for participant_number in sorted(participant_headers.keys()):
                participant_values = []
                for header in participant_headers[participant_number]:
                    value = row[header]
                    if pd.isna(value) or not str(value).strip():
                        continue

                    cleaned_header = re.sub(r'^(\s*\d+\.)+\s*', '', str(header)).strip()
                    cleaned_header = re.sub(
                        rf'^participant\s+{participant_number}\s+',
                        '',
                        cleaned_header,
                        flags=re.IGNORECASE
                    ).strip()
                    participant_values.append((cleaned_header, str(value).strip()))

                if not participant_values:
                    continue

                participant_heading = doc.add_paragraph()
                participant_heading_run = participant_heading.add_run(f"Participant {participant_number}")
                participant_heading_run.bold = True

                for cleaned_header, cleaned_value in participant_values:
                    bullet_paragraph = doc.add_paragraph(style='List Bullet')
                    header_run = bullet_paragraph.add_run(f"{cleaned_header}: ")
                    header_run.bold = True
                    bullet_paragraph.add_run(cleaned_value)

                if participant_number == 1:
                    session_values = []
                    for session_header in pre_symposium_headers:
                        session_value = row[session_header]
                        if pd.notna(session_value) and str(session_value).strip():
                            session_values.append(str(session_value).strip())

                    if session_values:
                        session_bullet = doc.add_paragraph(style='List Bullet')
                        session_header_run = session_bullet.add_run(
                            "Pre-symposium Virtual Learning Collaborative Sessions: "
                        )
                        session_header_run.bold = True
                        session_bullet.add_run(" | ".join(session_values))

            selected_sww_impacts_projects = []
            for project_header in sww_impacts_project_headers:
                project_value = row[project_header]
                if pd.isna(project_value):
                    continue

                value_text = str(project_value).strip().lower()
                if value_text == 'checked':
                    selected_sww_impacts_projects.append(sww_impacts_project_labels[project_header])

            if selected_sww_impacts_projects:
                p = doc.add_paragraph()
                header_run = p.add_run("Which SWW-IMPACTS project will you focus on?: ")
                header_run.bold = True
                p.add_run(" | ".join(selected_sww_impacts_projects))

            selected_success_story_types = []
            for success_header, success_label in success_story_headers.items():
                success_value = row[success_header]
                if pd.isna(success_value):
                    continue

                success_value_text = str(success_value).strip().lower()
                if success_value_text.startswith('yes'):
                    selected_success_story_types.append(success_label)

            p = doc.add_paragraph()
            header_run = p.add_run("Are you will to share success stories?: ")
            header_run.bold = True
            if selected_success_story_types:
                p.add_run(" | ".join(selected_success_story_types))
            else:
                p.add_run("No")

            # Write each column header with its corresponding value
            for header in headers:
                if header in likert_values_by_header:
                    if pd.notna(row[header]):
                        likert_values_by_header[header].append(int(row[header]))
                    continue

                if header in participant_header_set:
                    continue

                if header in sww_impacts_project_headers:
                    continue

                if header in success_story_headers:
                    continue

                if header in pre_symposium_headers:
                    continue

                if header in other_specify_headers:
                    other_value = row[header]
                    if pd.isna(other_value) or not str(other_value).strip():
                        continue

                value = row[header]
                p = doc.add_paragraph()
                # Add header in bold
                cleaned_header = re.sub(r'^(\s*\d+\.)+\s*', '', str(header)).strip()
                header_run = p.add_run(f"{cleaned_header}: ")
                header_run.bold = True
                # Add value in normal text
                p.add_run(str(value))

            # Add blank paragraph between rows
            doc.add_paragraph()

        # Add summary table for 1-5 scale columns
        if likert_headers:
            summary_heading = doc.add_paragraph()
            summary_heading_run = summary_heading.add_run("1-5 Scale Summary")
            summary_heading_run.bold = True

            summary_table = doc.add_table(rows=1, cols=6)
            summary_table.style = 'Table Grid'
            header_cells = summary_table.rows[0].cells
            header_cells[0].text = 'Column'
            header_cells[1].text = 'Mean'
            header_cells[2].text = 'Median'
            header_cells[3].text = 'Mode'
            header_cells[4].text = 'Range'
            header_cells[5].text = 'Count'

            for header in likert_headers:
                values = likert_values_by_header[header]
                cleaned_header = re.sub(r'^(\s*\d+\.)+\s*', '', str(header)).strip()

                row_cells = summary_table.add_row().cells
                row_cells[0].text = cleaned_header

                if not values:
                    row_cells[1].text = 'N/A'
                    row_cells[2].text = 'N/A'
                    row_cells[3].text = 'N/A'
                    row_cells[4].text = 'N/A'
                    row_cells[5].text = '0'
                    continue

                mean_value = statistics.mean(values)
                median_value = statistics.median(values)
                mode_values = sorted(statistics.multimode(values))
                min_value = min(values)
                max_value = max(values)

                row_cells[1].text = f"{mean_value:.2f}"
                row_cells[2].text = f"{median_value:.2f}"
                row_cells[3].text = ', '.join(str(mode_value) for mode_value in mode_values)
                row_cells[4].text = f"{min_value}-{max_value}"
                row_cells[5].text = str(len(values))

        # Save the document
        doc.save(output_file)

        print(f"Successfully converted {excel_file} to {output_file} using special parser")
        return output_file

    except FileNotFoundError:
        print(f"Error: File '{excel_file}' not found.")
        sys.exit(1)
    except Exception as e:
        print(f"Error: {e}")
        sys.exit(1)



def excel_to_text(excel_file: str, output_file: str = None) -> str:
    """
    Convert an Excel file to a formatted text document.
    Each row is written on separate lines with column headers preceding the data.

    Args:
        excel_file: Path to the Excel file
        output_file: Path to the output text file (optional)
    """
    try:
        # Read the input file
        df = read_input_file(excel_file)

        # If no output file specified, create one based on input filename
        if output_file is None:
            output_file = Path(excel_file).stem + "_output.txt"

        # Open the output file for writing
        with open(output_file, 'w', encoding='utf-8') as f:
            # Get column headers
            headers = df.columns.tolist()

            # Process each row
            for index, row in df.iterrows():
                f.write(f"Row {index + 1}:\n")
                f.write("-" * 50 + "\n")

                # Write each column header with its corresponding value
                for header in headers:
                    value = row[header]
                    f.write(f"{header}: {value}\n")

                f.write("\n")  # Add blank line between rows

        print(f"Successfully converted {excel_file} to {output_file}")
        return output_file

    except FileNotFoundError:
        print(f"Error: File '{excel_file}' not found.")
        sys.exit(1)
    except Exception as e:
        print(f"Error: {e}")
        sys.exit(1)


if __name__ == "__main__":
    if len(sys.argv) < 2:
        print("Usage: python excel_to_text.py <input_file> [output_file] [--format txt|docx] [--special]")
        print("Example: python excel_to_text.py data.xlsx output.txt")
        print("Example: python excel_to_text.py data.csv output.txt")
        print("Example: python excel_to_text.py data.xlsx output.docx --format docx")
        print("Example: python excel_to_text.py special.xlsx output.docx --format docx --special")
        sys.exit(1)

    excel_file = sys.argv[1]
    output_file = None
    output_format = 'txt'  # default format
    use_special_parser = False

    # Parse arguments
    for i, arg in enumerate(sys.argv[2:], start=2):
        if arg == '--format' and i + 1 < len(sys.argv):
            output_format = sys.argv[i + 1].lower()
        elif arg == '--special':
            use_special_parser = True
        elif not arg.startswith('--') and arg not in ['txt', 'docx']:
            output_file = arg

    # Auto-detect format from output file extension if provided
    if output_file:
        ext = Path(output_file).suffix.lower()
        if ext == '.docx':
            output_format = 'docx'
        elif ext == '.txt':
            output_format = 'txt'

    # Convert based on format
    if output_format == 'docx':
        if use_special_parser:
            specific_excel_to_word(excel_file, output_file)
        else:
            excel_to_word(excel_file, output_file)
    else:
        if use_special_parser:
            print("Warning: --special applies only to DOCX output. Using standard text conversion.")
        excel_to_text(excel_file, output_file)
